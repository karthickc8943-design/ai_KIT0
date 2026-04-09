"""
File management functions for AI directory
"""
import os
import shutil
from PIL import Image

try:
    from IPython.display import display, Markdown
    HAS_IPYTHON = True
except ImportError:
    HAS_IPYTHON = False

try:
    from .image import analyze_image
except ImportError:
    from image import analyze_image

# Constants
AI_FILES_DIR = os.path.expanduser("~/ai_files")
MAX_IMAGE_SIZE = 5 * 1024 * 1024  # 5MB

def setup_ai_directory():
    """Create AI files directory"""
    if not os.path.exists(AI_FILES_DIR):
        os.makedirs(AI_FILES_DIR)
        print(f"✅ Created AI files directory: {AI_FILES_DIR}")
    return AI_FILES_DIR

def list_ai_files():
    """List files in AI directory with markdown formatting"""
    setup_ai_directory()
    files = os.listdir(AI_FILES_DIR)

    if files:
        markdown = "## 📁 Files in AI Directory\n\n"
        markdown += "| # | Type | Filename | Size |\n"
        markdown += "|---|------|----------|------|\n"

        for i, file in enumerate(files, 1):
            path = os.path.join(AI_FILES_DIR, file)
            size = os.path.getsize(path) / 1024
            size_str = f"{size:.1f} KB" if size < 1024 else f"{size/1024:.1f} MB"
            ext = os.path.splitext(file)[1].lower()

            if ext in ('.png', '.jpg', '.jpeg'): ftype = "🖼️ Image"
            elif ext == '.pdf': ftype = "📑 PDF"
            elif ext == '.docx': ftype = "📝 Word"
            elif ext in ('.xlsx', '.xls', '.csv'): ftype = "📊 Spreadsheet"
            else: ftype = "📄 Text"
            
            markdown += f"| {i} | {ftype} | `{file}` | {size_str} |\n"

        if HAS_IPYTHON:
            display(Markdown(markdown))
        else:
            print(markdown)
    else:
        if HAS_IPYTHON:
            display(Markdown("📁 AI directory is empty"))
        else:
            print("📁 AI directory is empty")
    return files

def upload_to_ai(source_path):
    """Upload file to AI directory"""
    setup_ai_directory()

    if not os.path.exists(source_path):
        print(f"❌ Source file not found: {source_path}")
        return False

    filename = os.path.basename(source_path)
    dest_path = os.path.join(AI_FILES_DIR, filename)

    # Handle large images
    if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
        file_size = os.path.getsize(source_path)
        if file_size > MAX_IMAGE_SIZE:
            print(f"⚠️ Large image, resizing...")
            img = Image.open(source_path)
            img.thumbnail((1024, 1024))
            img.save(dest_path, quality=85)
            print(f"✅ Uploaded resized image")
            return dest_path

    shutil.copy2(source_path, dest_path)
    print(f"✅ Uploaded '{filename}'")
    return dest_path

def read_ai_file(filename):
    """Read a text file from AI directory"""
    setup_ai_directory()
    path = os.path.join(AI_FILES_DIR, filename)

    if not os.path.exists(path):
        print(f"❌ File not found: {filename}")
        return None

    with open(path, 'r') as f:
        content = f.read()

    if HAS_IPYTHON:
        display(Markdown(f"### 📄 {filename}\n\n```\n{content}\n```"))
    else:
        print(f"📄 {filename}\n\n{content}")
    return content

def process_ai_image(filename, prompt="Describe this image"):
    """Process an image from AI directory"""
    setup_ai_directory()
    path = os.path.join(AI_FILES_DIR, filename)

    if not os.path.exists(path):
        print(f"❌ Image not found: {filename}")
        return None

    result = analyze_image(path, prompt)
    if HAS_IPYTHON:
        display(Markdown(f"### 🖼️ {filename}\n\n{result}"))
    else:
        print(f"### 🖼️ {filename}\n\n{result}")
    return result

def delete_ai_file(filename):
    """Delete a file from AI directory"""
    setup_ai_directory()
    path = os.path.join(AI_FILES_DIR, filename)

    if os.path.exists(path):
        os.remove(path)
        print(f"✅ Deleted '{filename}'")
    else:
        print(f"❌ File not found: {filename}")

def read_ai_file_safe(filename):
    """
    Safely read a file - shows text files normally, warns for binary files
    """
    setup_ai_directory()
    path = os.path.join(AI_FILES_DIR, filename)

    if not os.path.exists(path):
        print(f"❌ File not found: {filename}")
        return None

    # Check if it's a text file
    text_extensions = ['.txt', '.md', '.csv', '.json', '.py', '.html', '.css', '.js', '.xml']
    ext = os.path.splitext(filename)[1].lower()

    if ext in text_extensions:
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        if HAS_IPYTHON:
            display(Markdown(f"### 📄 {filename}\n\n```\n{content[:2000]}{'...' if len(content)>2000 else ''}\n```"))
        else:
            print(f"### 📄 {filename}\n\n{content[:2000]}...")
        return content
    elif ext == '.pdf':
        print(f"📑 {filename} is a PDF. Use a PDF reader to view it.")
        return None
    else:
        print(f"📁 {filename} is a binary file ({ext})")
        return None

import PyPDF2
import docx
import pandas as pd
import pytesseract

def read_pdf(filename):
    """Read a PDF file and return text content"""
    path = os.path.join(AI_FILES_DIR, filename)
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = f"CONTENT OF {filename} (PDF):\n"
            for i, page in enumerate(pdf_reader.pages):
                text += f"\n[Page {i+1}]\n{page.extract_text()}\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def read_docx(filename):
    """Read a Word document and return text content"""
    path = os.path.join(AI_FILES_DIR, filename)
    try:
        doc = docx.Document(path)
        text = f"CONTENT OF {filename} (Word):\n"
        for i, para in enumerate(doc.paragraphs):
            text += f"{para.text}\n"
        return text
    except Exception as e:
        return f"Error reading DOCX: {e}"

def read_excel_csv(filename):
    """Read Excel or CSV file and return text content as markdown table"""
    path = os.path.join(AI_FILES_DIR, filename)
    try:
        if filename.endswith('.csv'):
            df = pd.read_csv(path)
        else:
            df = pd.read_excel(path)
        return f"CONTENT OF {filename} (Spreadsheet):\n\n{df.to_markdown(index=False)}"
    except Exception as e:
        return f"Error reading Spreadsheet: {e}"

def read_ocr_image(filename):
    """Extract text from an image using OCR"""
    path = os.path.join(AI_FILES_DIR, filename)
    try:
        text = pytesseract.image_to_string(Image.open(path))
        return f"TEXT EXTRACTED FROM IMAGE {filename}:\n\n{text}"
    except Exception as e:
        return f"Error performing OCR on image: {e}"

def extract_any_file(path):
    """Universal silent extraction from ANY file path (PDF, DOCX, CSV, Image, etc.)"""
    if not os.path.exists(path): return f"File not found: {path}"
    
    filename = os.path.basename(path)
    ext = os.path.splitext(filename)[1].lower()
    content = ""

    try:
        if ext in ('.png', '.jpg', '.jpeg'):
            content = pytesseract.image_to_string(Image.open(path))
            content = f"TEXT EXTRACTED FROM IMAGE {filename}:\n\n{content}"
        elif ext == '.pdf':
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = f"CONTENT OF {filename} (PDF):\n"
                for i, page in enumerate(pdf_reader.pages):
                    content += f"\n[Page {i+1}]\n{page.extract_text()}\n"
        elif ext == '.docx':
            doc = docx.Document(path)
            content = f"CONTENT OF {filename} (Word):\n"
            for para in doc.paragraphs:
                content += f"{para.text}\n"
        elif ext == '.doc':
            # Older binary word format
            return f"❌ Note: '{filename}' is a legacy .doc binary file. JARVIS requires modern .docx or .pdf for text extraction."
        elif ext in ('.xlsx', '.xls', '.csv'):
            if ext == '.csv': df = pd.read_csv(path)
            else: df = pd.read_excel(path)
            content = f"CONTENT OF {filename} (Spreadsheet):\n\n{df.to_markdown(index=False)}"
        else:
            # Text file fallback (ignore decoding errors to prevent strict binary crashes)
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        return content
    except Exception as e:
        return f"Error extracting {filename}: {e}"

def read_ai_file_auto(filename):
    """Unifed auto-detector for AI directory files"""
    setup_ai_directory()
    path = os.path.join(AI_FILES_DIR, filename)
    content = extract_any_file(path)

    if HAS_IPYTHON:
        display(Markdown(f"### 📄 {filename}\n\n{content[:3000]}..."))
    else:
        print(f"📄 {filename}\n\n{content[:3000]}...")
    return content

def extract_ai_file(filename):
    """Silent extraction for AI processing - compatible with ALL formats in AI dir"""
    path = os.path.join(AI_FILES_DIR, filename)
    return extract_any_file(path)

def display_table(data, headers=None):
    """Display data as a beautiful markdown table"""
    import pandas as pd
    try:
        if isinstance(data, list):
            if all(isinstance(row, dict) for row in data):
                df = pd.DataFrame(data)
            elif all(isinstance(row, list) for row in data):
                df = pd.DataFrame(data, columns=headers) if headers else pd.DataFrame(data)
            else:
                df = pd.DataFrame([data])
        else:
            df = pd.DataFrame(data)

        markdown_table = "## 📊 Data Table\n\n"
        markdown_table += df.to_markdown(index=False)

        if HAS_IPYTHON:
            display(Markdown(markdown_table))
        else:
            print(markdown_table)
        return df

    except Exception as e:
        print(f"❌ Error creating table: {e}")
        return None

def display_table_simple(data, title="Data Table"):
    """Simpler version - just display a list of lists as table"""
    if not data:
        print("❌ No data to display")
        return

    markdown = f"## {title}\n\n"
    if isinstance(data[0], list):
        markdown += "| " + " | ".join([f"Column {i+1}" for i in range(len(data[0]))]) + " |\n"
        markdown += "|" + "---|" * len(data[0]) + "\n"
        for row in data:
            markdown += "| " + " | ".join([str(cell) for cell in row]) + " |\n"
    elif isinstance(data[0], dict):
        headers = list(data[0].keys())
        markdown += "| " + " | ".join(headers) + " |\n"
        markdown += "|" + "---|" * len(headers) + "\n"
        for row in data:
            markdown += "| " + " | ".join([str(row.get(h, '')) for h in headers]) + " |\n"

    if HAS_IPYTHON:
        display(Markdown(markdown))
    else:
        print(markdown)
